#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
高级格式化脚本：根据模板文档的格式和样式格式化报告文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import sys
import re

def analyze_template_styles(template_doc):
    """分析模板文档的样式结构"""
    styles_info = {
        'heading1': {'font_name': '宋体', 'font_size': Pt(16), 'bold': True, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
        'heading2': {'font_name': '宋体', 'font_size': Pt(14), 'bold': True, 'alignment': None},
        'heading3': {'font_name': '宋体', 'font_size': Pt(12), 'bold': True, 'alignment': None},
        'normal': {'font_name': '宋体', 'font_size': Pt(12), 'bold': False, 'alignment': None},
    }
    
    # 从模板中提取实际样式
    for para in template_doc.paragraphs[:200]:
        style_name = para.style.name if para.style else "Normal"
        if para.runs:
            run = para.runs[0]
            if 'Heading 1' in style_name or '标题 1' in style_name:
                if run.font.name:
                    styles_info['heading1']['font_name'] = run.font.name
                if run.font.size:
                    styles_info['heading1']['font_size'] = run.font.size
            elif 'Heading 2' in style_name or '标题 2' in style_name:
                if run.font.name:
                    styles_info['heading2']['font_name'] = run.font.name
                if run.font.size:
                    styles_info['heading2']['font_size'] = run.font.size
    
    return styles_info

def detect_paragraph_type(text):
    """检测段落类型"""
    text = text.strip()
    
    # 空行
    if not text:
        return 'empty'
    
    # 分隔线
    if text == '---' or text.startswith('---'):
        return 'separator'
    
    # 一级标题：以# 开头或包含"第X章"
    if text.startswith('# ') or (text.startswith('第') and '章' in text and len(text) < 50):
        return 'heading1'
    
    # 二级标题：以## 开头或包含"X.X"格式
    if text.startswith('## ') or re.match(r'^\d+\.\d+\s+', text):
        return 'heading2'
    
    # 三级标题：以### 开头或包含"X.X.X"格式
    if text.startswith('### ') or re.match(r'^\d+\.\d+\.\d+\s+', text):
        return 'heading3'
    
    # 关键词段落
    if text.startswith('**关键词**') or text.startswith('关键词'):
        return 'keywords'
    
    # 目录项
    if text.startswith('- [') or text.startswith('* ['):
        return 'toc_item'
    
    # 说明文字
    if text.startswith('**说明**') or text.startswith('说明'):
        return 'note'
    
    # 普通段落
    return 'normal'

def apply_style_to_paragraph(para, style_info, text):
    """应用样式到段落"""
    # 清除现有内容
    para.clear()
    
    # 添加文本
    run = para.add_run(text)
    
    # 设置字体
    run.font.name = style_info['font_name']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), style_info['font_name'])
    
    # 设置字号
    if style_info['font_size']:
        run.font.size = style_info['font_size']
    
    # 设置加粗
    run.bold = style_info['bold']
    
    # 设置对齐
    if style_info['alignment']:
        para.alignment = style_info['alignment']
    
    # 设置段落格式
    para_format = para.paragraph_format
    para_format.space_after = Pt(6)  # 段后间距
    para_format.first_line_indent = Pt(24) if style_info.get('indent', False) else Pt(0)  # 首行缩进

def format_document(source_doc, template_doc, output_path):
    """格式化文档"""
    # 分析模板样式
    styles_info = analyze_template_styles(template_doc)
    
    # 创建新文档
    new_doc = Document()
    
    # 设置默认字体
    styles = new_doc.styles
    normal_style = styles['Normal']
    font = normal_style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 处理源文档
    for para in source_doc.paragraphs:
        text = para.text.strip()
        para_type = detect_paragraph_type(text)
        
        if para_type == 'empty':
            new_doc.add_paragraph()
            continue
        
        if para_type == 'separator':
            continue
        
        # 创建新段落
        new_para = new_doc.add_paragraph()
        
        # 根据类型应用样式
        if para_type == 'heading1':
            # 移除Markdown标记
            clean_text = re.sub(r'^#+\s*', '', text)
            clean_text = clean_text.strip()
            apply_style_to_paragraph(new_para, styles_info['heading1'], clean_text)
            new_para.style = 'Heading 1'
        
        elif para_type == 'heading2':
            # 移除Markdown标记
            clean_text = re.sub(r'^#+\s*', '', text)
            clean_text = clean_text.strip()
            apply_style_to_paragraph(new_para, styles_info['heading2'], clean_text)
            new_para.style = 'Heading 2'
        
        elif para_type == 'heading3':
            # 移除Markdown标记
            clean_text = re.sub(r'^#+\s*', '', text)
            clean_text = clean_text.strip()
            apply_style_to_paragraph(new_para, styles_info['heading3'], clean_text)
            new_para.style = 'Heading 3'
        
        elif para_type == 'keywords':
            # 关键词段落，加粗
            clean_text = text.replace('**', '')
            apply_style_to_paragraph(new_para, {**styles_info['normal'], 'bold': True}, clean_text)
        
        elif para_type == 'toc_item':
            # 目录项，保持原样但格式化
            clean_text = re.sub(r'^[-*]\s*\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            apply_style_to_paragraph(new_para, styles_info['normal'], clean_text)
            new_para.paragraph_format.left_indent = Inches(0.5)
        
        elif para_type == 'note':
            # 说明文字，可以设置为小字号或斜体
            clean_text = text.replace('**', '')
            apply_style_to_paragraph(new_para, {**styles_info['normal'], 'font_size': Pt(10)}, clean_text)
        
        else:
            # 普通段落
            # 处理Markdown格式
            clean_text = text
            # 移除加粗标记但保留文本
            clean_text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', clean_text)
            # 移除链接标记
            clean_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_text)
            
            apply_style_to_paragraph(new_para, {**styles_info['normal'], 'indent': True}, clean_text)
    
    # 保存文档
    new_doc.save(output_path)
    print(f"格式化完成，已保存到: {output_path}")

def main():
    """主函数"""
    template_path = "23050342008_高榆展_基于B2B的可行性分析报告智能体.docx"
    source_path = "报告_01_封面和摘要.docx"
    output_path = "报告_01_封面和摘要_格式化.docx"
    
    if not os.path.exists(template_path):
        print(f"错误: 模板文件不存在: {template_path}")
        return
    
    if not os.path.exists(source_path):
        print(f"错误: 源文件不存在: {source_path}")
        return
    
    print("正在读取模板文档...")
    template_doc = Document(template_path)
    
    print("正在读取源文档...")
    source_doc = Document(source_path)
    
    print("正在分析模板样式...")
    print("正在应用格式...")
    format_document(source_doc, template_doc, output_path)
    
    print("完成！")

if __name__ == "__main__":
    main()

