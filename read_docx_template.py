#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""读取docx模板文件，提取结构和内容"""

import docx
import os
import sys

# 获取当前目录下的docx文件
current_dir = os.path.dirname(os.path.abspath(__file__))
docx_files = [f for f in os.listdir(current_dir) if f.endswith('.docx')]

if not docx_files:
    print("未找到docx文件")
    sys.exit(1)

docx_file = docx_files[0]
print(f"正在读取文件: {docx_file}")

try:
    doc = docx.Document(docx_file)
    
    print(f"\n=== 文件基本信息 ===")
    print(f"段落总数: {len(doc.paragraphs)}")
    print(f"表格总数: {len(doc.tables)}")
    
    print(f"\n=== 文档结构分析 ===")
    
    # 分析段落结构
    chapter_count = 0
    section_count = 0
    subsection_count = 0
    
    structure = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        # 检测章节标题
        style_name = para.style.name if para.style else "Normal"
        
        # 判断标题级别
        if 'Heading 1' in style_name or (len(text) < 50 and ('章' in text or '第' in text and '章' in text)):
            chapter_count += 1
            structure.append(f"章节 {chapter_count}: {text}")
        elif 'Heading 2' in style_name or (len(text) < 30 and (text.startswith('第') or text.startswith('1.') or text.startswith('2.'))):
            section_count += 1
            structure.append(f"  节 {section_count}: {text}")
        elif 'Heading 3' in style_name or (len(text) < 20 and (text.startswith('1.1') or text.startswith('2.1'))):
            subsection_count += 1
    
    print(f"章节数: {chapter_count}")
    print(f"节数: {section_count}")
    print(f"小节数: {subsection_count}")
    
    print(f"\n=== 前100个段落预览 ===")
    for i, para in enumerate(doc.paragraphs[:100]):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "Normal"
            print(f"{i+1:3d}. [{style:15s}] {text[:80]}")
    
    print(f"\n=== 表格信息 ===")
    for i, table in enumerate(doc.tables):
        print(f"表格 {i+1}: {len(table.rows)}行 x {len(table.columns)}列")
        if i < 3:  # 只显示前3个表格的内容
            print("  内容预览:")
            for row_idx, row in enumerate(table.rows[:3]):
                row_data = [cell.text.strip()[:20] for cell in row.cells]
                print(f"    行{row_idx+1}: {' | '.join(row_data)}")
    
    # 统计字数
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    print(f"\n=== 内容统计 ===")
    print(f"总字符数: {total_chars:,}")
    print(f"总段落数: {len([p for p in doc.paragraphs if p.text.strip()])}")
    
except Exception as e:
    print(f"读取文件时出错: {e}")
    import traceback
    traceback.print_exc()

