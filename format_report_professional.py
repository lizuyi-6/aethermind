#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
专业格式化脚本：精确复制模板文档的格式和样式
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys
import re

def copy_style_properties(source_style, target_style):
    """复制样式属性"""
    try:
        # 复制字体
        if hasattr(source_style, 'font') and source_style.font:
            if source_style.font.name:
                target_style.font.name = source_style.font.name
            if source_style.font.size:
                target_style.font.size = source_style.font.size
            if source_style.font.bold is not None:
                target_style.font.bold = source_style.font.bold
        
        # 复制段落格式
        if hasattr(source_style, 'paragraph_format') and source_style.paragraph_format:
            source_pf = source_style.paragraph_format
            target_pf = target_style.paragraph_format
            
            if source_pf.alignment:
                target_pf.alignment = source_pf.alignment
            if source_pf.space_before:
                target_pf.space_before = source_pf.space_before
            if source_pf.space_after:
                target_pf.space_after = source_pf.space_after
            if source_pf.first_line_indent:
                target_pf.first_line_indent = source_pf.first_line_indent
            if source_pf.line_spacing:
                target_pf.line_spacing = source_pf.line_spacing
    except Exception as e:
        print(f"复制样式时出错: {e}")

def get_template_style_mapping(template_doc):
    """获取模板的样式映射"""
    style_mapping = {}
    
    # 分析模板中的样式使用
    for para in template_doc.paragraphs[:300]:
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()
        
        if not text:
            continue
        
        # 根据文本特征判断样式类型
        if 'Heading 1' in style_name or '标题 1' in style_name:
            style_mapping['heading1'] = style_name
        elif 'Heading 2' in style_name or '标题 2' in style_name:
            style_mapping['heading2'] = style_name
        elif 'Heading 3' in style_name or '标题 3' in style_name:
            style_mapping['heading3'] = style_name
        
        # 检查是否是章节标题
        if text.startswith('第') and '章' in text:
            if 'heading1' not in style_mapping:
                style_mapping['heading1'] = style_name
        elif re.match(r'^\d+\.\d+\s+', text):
            if 'heading2' not in style_mapping:
                style_mapping['heading2'] = style_name
        elif re.match(r'^\d+\.\d+\.\d+\s+', text):
            if 'heading3' not in style_mapping:
                style_mapping['heading3'] = style_name
    
    return style_mapping

def set_chinese_font(run, font_name='宋体'):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def format_paragraph_by_template(new_para, text, template_style, template_doc):
    """根据模板样式格式化段落"""
    # 清除现有内容
    new_para.clear()
    
    # 添加文本
    run = new_para.add_run(text)
    
    # 应用模板样式
    if template_style:
        new_para.style = template_style.name if hasattr(template_style, 'name') else template_style
        
        # 复制样式属性
        if hasattr(template_style, 'font') and template_style.font:
            if template_style.font.name:
                set_chinese_font(run, template_style.font.name)
            else:
                set_chinese_font(run, '宋体')
            
            if template_style.font.size:
                run.font.size = template_style.font.size
            else:
                run.font.size = Pt(12)
            
            if template_style.font.bold is not None:
                run.bold = template_style.font.bold
        else:
            set_chinese_font(run, '宋体')
            run.font.size = Pt(12)
    else:
        set_chinese_font(run, '宋体')
        run.font.size = Pt(12)
    
    # 设置段落格式
    if template_style and hasattr(template_style, 'paragraph_format'):
        source_pf = template_style.paragraph_format
        target_pf = new_para.paragraph_format
        
        if source_pf.alignment:
            target_pf.alignment = source_pf.alignment
        if source_pf.space_before:
            target_pf.space_before = source_pf.space_before
        if source_pf.space_after:
            target_pf.space_after = source_pf.space_after
        if source_pf.first_line_indent:
            target_pf.first_line_indent = source_pf.first_line_indent

def format_document_professional(source_doc, template_doc, output_path):
    """专业格式化文档"""
    # 获取模板样式映射
    style_mapping = get_template_style_mapping(template_doc)
    
    # 创建新文档
    new_doc = Document()
    
    # 复制模板的样式到新文档
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Normal', 'Title']:
        try:
            template_style = template_doc.styles[style_name]
            new_style = new_doc.styles[style_name]
            copy_style_properties(template_style, new_style)
        except:
            pass
    
    # 设置默认字体
    normal_style = new_doc.styles['Normal']
    normal_style.font.name = '宋体'
    normal_style.font.size = Pt(12)
    
    # 处理源文档
    for para in source_doc.paragraphs:
        text = para.text.strip()
        
        # 空行
        if not text:
            new_doc.add_paragraph()
            continue
        
        # 分隔线
        if text == '---' or text.startswith('---'):
            continue
        
        # 创建新段落
        new_para = new_doc.add_paragraph()
        
        # 判断段落类型
        clean_text = text
        
        # 一级标题
        if text.startswith('# ') or (text.startswith('第') and '章' in text and len(text) < 50):
            clean_text = re.sub(r'^#+\s*', '', text).strip()
            try:
                template_style = template_doc.styles['Heading 1']
            except:
                template_style = new_doc.styles['Heading 1']
            format_paragraph_by_template(new_para, clean_text, template_style, template_doc)
        
        # 二级标题
        elif text.startswith('## ') or re.match(r'^\d+\.\d+\s+', text):
            clean_text = re.sub(r'^#+\s*', '', text).strip()
            try:
                template_style = template_doc.styles['Heading 2']
            except:
                template_style = new_doc.styles['Heading 2']
            format_paragraph_by_template(new_para, clean_text, template_style, template_doc)
        
        # 三级标题
        elif text.startswith('### ') or re.match(r'^\d+\.\d+\.\d+\s+', text):
            clean_text = re.sub(r'^#+\s*', '', text).strip()
            try:
                template_style = template_doc.styles['Heading 3']
            except:
                template_style = new_doc.styles['Heading 3']
            format_paragraph_by_template(new_para, clean_text, template_style, template_doc)
        
        # 关键词
        elif text.startswith('**关键词**') or text.startswith('关键词'):
            clean_text = text.replace('**', '').strip()
            run = new_para.add_run(clean_text)
            set_chinese_font(run, '宋体')
            run.font.size = Pt(12)
            run.bold = True
        
        # 目录项
        elif text.startswith('- [') or text.startswith('* ['):
            clean_text = re.sub(r'^[-*]\s*\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            run = new_para.add_run(clean_text)
            set_chinese_font(run, '宋体')
            run.font.size = Pt(12)
            new_para.paragraph_format.left_indent = Cm(1)
        
        # 普通段落
        else:
            # 清理Markdown格式
            clean_text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', clean_text)  # 移除加粗
            clean_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_text)  # 移除链接
            
            try:
                template_style = template_doc.styles['Normal']
            except:
                template_style = new_doc.styles['Normal']
            format_paragraph_by_template(new_para, clean_text, template_style, template_doc)
            
            # 普通段落首行缩进
            new_para.paragraph_format.first_line_indent = Cm(0.74)  # 2字符缩进
    
    # 保存文档
    new_doc.save(output_path)
    print(f"格式化完成，已保存到: {output_path}")

def main():
    """主函数"""
    template_path = "23050342008_高榆展_基于B2B的可行性分析报告智能体.docx"
    source_path = "报告_01_封面和摘要.docx"
    output_path = "报告_01_封面和摘要_格式化.docx"
    
    if not os.path.exists(template_path):
        print(f"错误: 模板文件不存在: {template_path}")
        return
    
    if not os.path.exists(source_path):
        print(f"错误: 源文件不存在: {source_path}")
        return
    
    print("正在读取模板文档...")
    template_doc = Document(template_path)
    
    print("正在读取源文档...")
    source_doc = Document(source_path)
    
    print("正在分析模板样式...")
    print("正在应用格式...")
    format_document_professional(source_doc, template_doc, output_path)
    
    print("完成！输出文件: " + output_path)

if __name__ == "__main__":
    main()

