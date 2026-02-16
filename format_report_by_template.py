#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
根据模板文档的格式和样式格式化报告文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import sys

def get_template_styles(template_doc):
    """从模板文档中提取样式信息"""
    styles = {}
    
    # 提取段落样式
    for para in template_doc.paragraphs[:100]:  # 只检查前100个段落
        style_name = para.style.name if para.style else "Normal"
        if style_name not in styles:
            styles[style_name] = {
                'name': style_name,
                'font_name': None,
                'font_size': None,
                'bold': False,
                'alignment': None
            }
            
            # 获取字体信息
            if para.runs:
                run = para.runs[0]
                if run.font.name:
                    styles[style_name]['font_name'] = run.font.name
                if run.font.size:
                    styles[style_name]['font_size'] = run.font.size
                styles[style_name]['bold'] = run.bold if run.bold is not None else False
            
            # 获取对齐方式
            if para.alignment:
                styles[style_name]['alignment'] = para.alignment
    
    return styles

def apply_template_formatting(source_doc, template_doc, output_path):
    """将源文档按照模板格式格式化"""
    
    # 创建新文档
    new_doc = Document()
    
    # 复制模板的样式
    for style in template_doc.styles:
        try:
            new_style = new_doc.styles[style.name]
        except KeyError:
            new_style = new_doc.styles.add_style(style.name, style.type)
        
        # 复制字体设置
        if hasattr(style, 'font'):
            new_style.font.name = style.font.name
            if style.font.size:
                new_style.font.size = style.font.size
    
    # 处理源文档的每个段落
    for para in source_doc.paragraphs:
        text = para.text.strip()
        if not text:
            # 保留空行
            new_doc.add_paragraph()
            continue
        
        # 判断段落类型并应用相应样式
        new_para = None
        
        # 检测标题
        if text.startswith('# '):
            # 一级标题
            new_para = new_doc.add_paragraph(text[2:].strip())
            new_para.style = 'Heading 1'
        elif text.startswith('## '):
            # 二级标题
            new_para = new_doc.add_paragraph(text[3:].strip())
            new_para.style = 'Heading 2'
        elif text.startswith('### '):
            # 三级标题
            new_para = new_doc.add_paragraph(text[4:].strip())
            new_para.style = 'Heading 3'
        elif text.startswith('第') and ('章' in text or '节' in text):
            # 章节标题
            if '章' in text:
                new_para = new_doc.add_paragraph(text)
                new_para.style = 'Heading 1'
            elif '节' in text:
                new_para = new_doc.add_paragraph(text)
                new_para.style = 'Heading 2'
        elif text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) and len(text) < 100:
            # 可能是小节标题
            new_para = new_doc.add_paragraph(text)
            new_para.style = 'Heading 3'
        elif text.startswith('**') and text.endswith('**'):
            # 加粗文本
            new_para = new_doc.add_paragraph(text.strip('*'))
            new_para.runs[0].bold = True
        elif text == '---':
            # 分隔线，跳过
            continue
        else:
            # 普通段落
            new_para = new_doc.add_paragraph(text)
            new_para.style = 'Normal'
        
        # 设置中文字体
        if new_para:
            for run in new_para.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if not run.font.size:
                    run.font.size = Pt(12)
    
    # 保存新文档
    new_doc.save(output_path)
    print(f"格式化完成，已保存到: {output_path}")

def format_report():
    """主函数"""
    template_path = "23050342008_高榆展_基于B2B的可行性分析报告智能体.docx"
    source_path = "报告_01_封面和摘要.docx"
    output_path = "报告_01_封面和摘要_格式化.docx"
    
    if not os.path.exists(template_path):
        print(f"错误: 模板文件不存在: {template_path}")
        return
    
    if not os.path.exists(source_path):
        print(f"错误: 源文件不存在: {source_path}")
        return
    
    print("正在读取模板文档...")
    template_doc = Document(template_path)
    
    print("正在读取源文档...")
    source_doc = Document(source_path)
    
    print("正在应用模板格式...")
    apply_template_formatting(source_doc, template_doc, output_path)
    
    print("完成！")

if __name__ == "__main__":
    format_report()

