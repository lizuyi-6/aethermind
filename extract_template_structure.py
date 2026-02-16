#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""提取docx模板的完整结构"""

import docx
import os
import re

docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]
if not docx_files:
    print("未找到docx文件")
    exit(1)

doc = docx.Document(docx_files[0])
print(f"=== 完整章节结构 ===\n")

current_chapter = None
current_section = None
chapter_num = 0
section_num = 0

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    # 检测章节标题（第X章）
    if re.match(r'^第[一二三四五六七八九十\d]+章', text) or ('章' in text and len(text) < 50 and '第' in text):
        chapter_num += 1
        current_chapter = text
        current_section = None
        section_num = 0
        print(f"\n{'='*60}")
        print(f"第{chapter_num}章: {text}")
        print(f"{'='*60}")
    # 检测节标题（X.X）
    elif re.match(r'^\d+\.\d+', text) or (len(text) < 50 and text[0].isdigit() and '.' in text[:5]):
        section_num += 1
        current_section = text
        print(f"\n  {section_num}. {text}")
    # 显示部分内容
    elif current_chapter and len(text) > 20:
        # 只显示前3段内容作为示例
        if section_num <= 3:
            preview = text[:80].replace('\n', ' ')
            print(f"    - {preview}...")

print(f"\n\n=== 统计信息 ===")
print(f"总段落数: {len(doc.paragraphs)}")
print(f"非空段落数: {len([p for p in doc.paragraphs if p.text.strip()])}")
print(f"总字符数: {sum(len(p.text) for p in doc.paragraphs):,}")

