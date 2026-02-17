from __future__ import annotations

from pathlib import Path

from docx import Document


def main() -> None:
    here = Path(".")
    docx_files = sorted(here.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    # 只抽查最近生成的前几个 docx，避免扫描大量历史文件
    sample = docx_files[:5]
    if not sample:
        print("未找到 .docx 文件。")
        return

    for p in sample:
        try:
            doc = Document(str(p))
        except Exception as e:
            print(f"\n=== {p.name} ===")
            print("读取失败：", e)
            continue

        paras = [t.strip() for t in (pp.text for pp in doc.paragraphs) if t and t.strip()]
        bullet_cnt = sum(1 for pp in doc.paragraphs if getattr(pp.style, "name", "") == "List Bullet")
        print(f"\n=== {p.name} ===")
        joined = "\n".join(paras)
        if "WaytoAGI" in "\n".join(paras):
            print("[包含关键词] WaytoAGI")
        if "Vibe Coding" in joined:
            print("[包含关键词] Vibe Coding")
        if "AI 赋能" in joined or "AI赋能" in joined:
            print("[包含关键词] AI赋能")
        if bullet_cnt:
            print(f"[样式统计] List Bullet 段落数: {bullet_cnt}")
        for line in paras[:12]:
            print(line)
        print("... 段落数:", len(paras))


if __name__ == "__main__":
    main()


