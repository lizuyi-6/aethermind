from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TODAY_CN = datetime.now().strftime("%Y年%m月%d日")


def _set_default_font(doc: Document, font_name: str = "宋体", font_size_pt: int = 12) -> None:
    """
    为 Word 设置默认字体（含中文字体映射），尽量避免在不同电脑上乱码/字体不统一。
    """
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    # 设置东亚字体
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _title(doc: Document, text: str, subtitle: Optional[str] = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(14)
        r2.font.name = "宋体"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _meta_block(
    doc: Document,
    school: str = "（可填写学校/学院）",
    major: str = "（可填写专业）",
    class_name: str = "（可填写班级）",
    name: str = "高榆展",
    student_id: str = "23050342008",
) -> None:
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("学    校 / 学    院", school),
        ("专    业", major),
        ("班    级", class_name),
        ("姓    名 / 学    号", f"{name} / {student_id}"),
        ("日    期", TODAY_CN),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
    doc.add_paragraph()


def _h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullets(doc: Document, items: Iterable[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def _numbered(doc: Document, items: Iterable[str]) -> None:
    """
    数字编号列表：每次调用都从 1 重新开始（避免整篇文档共用一组连续编号）。
    """
    num_id = _add_restart_numbering_for_list_number(doc)
    for it in items:
        p = doc.add_paragraph(it, style="List Number")
        _force_paragraph_num_id(p, num_id, ilvl=0)


def _get_style_num_id(doc: Document, style_name: str) -> Optional[int]:
    """
    从样式定义中读取该段落样式绑定的 numId（若存在）。
    """
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None
    el = style.element
    vals = el.xpath(".//w:pPr/w:numPr/w:numId/@w:val")
    if not vals:
        return None
    try:
        return int(vals[0])
    except Exception:
        return None


def _next_num_id(numbering) -> int:
    vals = numbering.xpath(".//w:num/@w:numId")
    existing = []
    for v in vals:
        try:
            existing.append(int(v))
        except Exception:
            pass
    return (max(existing) + 1) if existing else 1


def _abstract_num_id_for_num_id(numbering, num_id: int) -> Optional[int]:
    vals = numbering.xpath(f".//w:num[@w:numId='{num_id}']/w:abstractNumId/@w:val")
    if not vals:
        return None
    try:
        return int(vals[0])
    except Exception:
        return None


def _add_restart_numbering_for_list_number(doc: Document) -> int:
    """
    为“List Number”创建一个新的 numId，并指定 lvl0 从 1 开始。
    """
    numbering_part = doc.part.numbering_part
    # python-docx 不同版本的属性名略有差异：此处用 element（底层 numbering.xml 根节点）
    numbering = numbering_part.element

    base_num_id = _get_style_num_id(doc, "List Number")
    if base_num_id is None:
        # 兜底：若样式没有绑定 numId，则直接使用样式默认；但我们依旧尝试创建一个新 num
        base_num_id = 1

    abstract_id = _abstract_num_id_for_num_id(numbering, base_num_id)
    if abstract_id is None:
        # 再兜底：取第一个 abstractNum
        abs_vals = numbering.xpath(".//w:abstractNum/@w:abstractNumId")
        abstract_id = int(abs_vals[0]) if abs_vals else 0

    new_num_id = _next_num_id(numbering)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))

    abs_el = OxmlElement("w:abstractNumId")
    abs_el.set(qn("w:val"), str(abstract_id))
    num.append(abs_el)

    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)

    numbering.append(num)
    return new_num_id


def _force_paragraph_num_id(paragraph, num_id: int, ilvl: int = 0) -> None:
    """
    强制把段落绑定到指定 numId/ilvl（从而实现每个列表独立编号）。
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)

    ilvl_el = numPr.find(qn("w:ilvl"))
    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        numPr.append(ilvl_el)
    ilvl_el.set(qn("w:val"), str(ilvl))

    numId_el = numPr.find(qn("w:numId"))
    if numId_el is None:
        numId_el = OxmlElement("w:numId")
        numPr.append(numId_el)
    numId_el.set(qn("w:val"), str(num_id))


def _code_block(doc: Document, code: str) -> None:
    """
    伪代码/代码段：使用等宽字体（英文 Courier New + 中文等宽不强制）。
    """
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(10)


def build_doc_1_understanding(path: str) -> None:
    doc = Document()
    _set_default_font(doc, "宋体", 12)

    _title(doc, "人工智能：个人理解与学习感悟", "（作业 1：对智能与人工智能的认识）")
    _meta_block(doc)

    _h1(doc, "一、我对“智能”的理解")
    _para(
        doc,
        "我理解的“智能”不是单一能力，而是一组面向目标的综合能力：感知信息、形成表示、做出推理与决策、在反馈中学习与改进，并在环境变化时保持一定的鲁棒性。"
        "人类智能之所以显得“强”，不仅因为会算、会记，更重要的是能在不完整信息下做取舍、能迁移经验、能用语言与工具扩展自身能力。"
        "因此，“智能”应当同时包含：知识与经验（记忆）、方法与策略（推理/规划）、以及自我修正（学习）。"
    )
    _numbered(
        doc,
        [
            "目标性：智能体的行为围绕目标展开，目标可以显式（完成任务）也可以隐式（维持安全/节能）。",
            "适应性：面对新情况能调整策略，而不是只能执行固定流程。",
            "可解释性（在多数场景下）：能给出“为什么这样做”的理由，便于协作与纠错。",
        ],
    )

    _h1(doc, "二、我对“人工智能”的理解")
    _para(
        doc,
        "我认为人工智能（AI）是用算法与系统工程的方法，让计算机在特定任务上表现出类似智能的能力。"
        "从历史上看，AI 大致经历了“符号主义（规则/逻辑）—统计学习（概率/优化）—深度学习（表示学习）—大模型与智能体（通用能力+工具使用）”的演进。"
        "我更认可把 AI 看成一套“可计算的智能方法论”：用数据与规则描述世界，用模型进行推断与生成，用反馈持续改进。"
    )
    _numbered(
        doc,
        [
            "弱人工智能：在特定任务上表现优秀（如推荐、识别、翻译、检索），当前主流落在这一层面。",
            "强人工智能：具备跨领域泛化、长期自主学习与稳健规划的通用智能，目前仍是研究目标与工程挑战。",
            "大模型带来的变化：把“能力”从单任务模型提升为可迁移的通用表征，并通过提示词/工具调用形成“可编排的能力”。",
        ],
    )

    _h1(doc, "三、我感兴趣的研究方向（及原因）")
    _h2(doc, "1. 大模型与智能体（Agent）")
    _para(
        doc,
        "原因：大模型具备较强的语言理解与生成能力，而智能体强调“目标-计划-执行-反思”的闭环。两者结合后，AI 不再只是回答问题，而是能调用工具、分解任务、做多步推理与迭代改进。"
    )
    _numbered(doc, ["提示工程与对话策略（Prompt/Memory）", "工具调用与函数执行（Tool Use）", "反思与自我纠错（Reflection）", "多智能体协作（Multi-Agent）"])

    _h2(doc, "2. 多模态学习与应用")
    _para(
        doc,
        "原因：生活中信息不仅是文本，还包括图像、音频、视频。多模态让 AI 能“看懂+听懂+会说”，更接近真实场景的输入输出形态。"
    )
    _numbered(doc, ["图文检索与理解", "视觉问答（VQA）", "多模态内容生成与安全过滤"])

    _h2(doc, "3. 可靠性、安全与伦理")
    _para(
        doc,
        "原因：AI 进入学习与生活后，正确性、偏见、隐私、版权、可追责等问题变得重要。技术上需要评测与对齐，使用上需要规范与边界。"
    )
    _numbered(doc, ["幻觉与事实一致性", "数据隐私与安全", "公平性与偏见治理", "可解释与可审计"])

    _h1(doc, "四、我如何学习这方面内容（学习路径与方法）")
    _para(
        doc,
        "我的学习方法偏向“先建立结构，再做小实验”。先用课程与书籍搭建知识框架，再用可复现实验把概念变成技能。"
        "我把学习拆成四条主线：数学基础、编程能力、算法与论文、工程实践。"
    )
    _numbered(
        doc,
        [
            "数学与基础：概率统计（贝叶斯/分布）、线性代数（向量/矩阵/特征值）、微积分与优化（梯度下降/正则化）。",
            "编程与工具：Python、Numpy/Pandas、可视化、Git；深度学习框架（PyTorch 或 TensorFlow）。",
            "算法与论文：从经典算法开始（贝叶斯、搜索、进化算法、神经网络），再读综述与代表性论文，形成“问题-方法-指标-局限”的思维模板。",
            "工程与项目：做小项目（分类/检索/生成），记录数据处理、训练、评估、上线/部署的完整流程。",
        ],
    )
    _numbered(
        doc,
        [
            "做笔记：把每个专题整理成“定义—公式—伪代码—实验—坑点—扩展阅读”。",
            "做复盘：对每次实验记录参数、结果、失败原因与改进方向。",
            "做对比：同一任务用不同方法做 baseline，对比优缺点，避免只会“套模型”。",
        ],
    )

    _h1(doc, "五、AI 对生活与学习的影响（我的感悟）")
    _h2(doc, "1. 对学习效率的提升")
    _numbered(
        doc,
        [
            "查资料更快：AI 能把分散的信息整理成结构化要点，降低入门门槛。",
            "写作与表达更顺：用于大纲、措辞优化、示例生成，但最终观点仍需自己把关。",
            "编程更高效：代码补全、报错定位、重构建议显著节省时间。",
        ],
    )
    _h2(doc, "2. 对学习方式的改变")
    _numbered(
        doc,
        [
            "从“记忆式学习”转向“问题驱动学习”：先提出问题，再用 AI 辅助探索路径。",
            "从“单一资料”转向“多源交叉验证”：AI 给答案后必须回到教材/论文/官方文档核对。",
        ],
    )
    _h2(doc, "3. 风险与自我要求")
    _numbered(
        doc,
        [
            "幻觉风险：AI 可能编造引用或细节，因此引用与结论必须可追溯。",
            "依赖风险：如果完全让 AI 代做，会导致思考能力退化；我会把 AI 当作“助教/搭档”，而不是“替代者”。",
            "伦理与合规：尊重隐私与版权，作业与论文中明确标注 AI 辅助范围，避免不当使用。",
        ],
    )
    _para(doc, "总体而言，我把 AI 当作一种“放大器”：它能放大效率，也会放大错误；关键在于人的判断、验证与责任。")

    _h1(doc, "六、结语")
    _para(
        doc,
        "未来我希望在“大模型+智能体+可靠性”的交叉方向持续学习，通过更多可复现实验与小项目，把知识沉淀成可迁移的能力。"
    )

    _h1(doc, "七、我对 Vibe Coding 的理解与实践")
    _para(
        doc,
        "我理解的 Vibe Coding 是一种“与 AI 搭档式编程”的学习与开发方式：先用自然语言把需求、约束与验收标准说清楚，"
        "再让 AI 生成初版方案与代码骨架，随后由人进行审阅、验证、补测试与迭代。它强调高频反馈与快速成型，但并不等于把思考完全交给 AI。"
    )
    _numbered(
        doc,
        [
            "我的常用流程：明确目标与边界 → 让 AI 输出方案/目录结构 → 小步实现并频繁运行 → 让 AI 辅助做代码审查与重构 → 自己补齐测试与文档。",
            "适合的任务：原型搭建、重复性脚手架、接口封装、报错定位、代码重构与注释/README 生成。",
            "主要风险：AI 可能写出“能跑但不对”的逻辑（或引入安全隐患）。对策是：写验收用例、加单元测试、关键逻辑手工复核、对外部依赖做版本锁定。",
            "对学习的帮助：把精力从语法细节转移到架构设计、边界条件、评估指标与工程规范上，学习效率更高。",
        ],
    )
    _para(
        doc,
        "总体上，我把 Vibe Coding 当作一种“加速试错”的方法：用 AI 提升效率，但以验证、复盘与责任来保证质量。"
    )

    _h1(doc, "八、AI 赋能：对学习与生活的具体帮助（我的做法）")
    _para(
        doc,
        "我理解的“AI 赋能”不是简单地把任务交给 AI 完成，而是把 AI 当作第二大脑与助教：用它加速检索、梳理、生成初稿与定位问题，"
        "同时用自己的验证与反思保证正确性、可解释性与合规性。"
    )
    _h2(doc, "1. 学习赋能：从“看懂”到“能做”的加速")
    _numbered(
        doc,
        [
            "快速建立知识框架：先让 AI 给出概念地图与学习路线，再对照教材/课程逐条补全，避免碎片化学习。",
            "辅助理解难点：对公式/推导/概念用“换一种讲法 + 举例 + 反例”的方式提问，直到能用自己的话复述。",
            "生成练习与自测：让 AI 出题（含答案与解析），再用错题本记录易错点，形成闭环。",
            "做学习复盘：把每周学习内容总结为“学了什么—做了什么—效果如何—下周改进”，提升长期稳定性。",
        ],
    )
    _h2(doc, "2. 写作与表达赋能：把时间用在观点与结构上")
    _numbered(
        doc,
        [
            "先自己定观点与大纲：AI 负责润色与扩写，但核心论点、逻辑顺序和结论由我决定。",
            "格式化输出：让 AI 按要求输出摘要、要点、对比表格与结论段，减少排版与措辞成本。",
            "引用与事实核对：凡是涉及数据、来源与结论的内容，必须回到原始文献/官方文档核验后再写入。",
        ],
    )
    _h2(doc, "3. 编程与工程赋能：提高开发效率但不牺牲质量")
    _numbered(
        doc,
        [
            "需求澄清与接口设计：用 AI 帮我把需求转成接口、数据结构与边界条件清单，减少返工。",
            "生成脚手架与样例：用 AI 生成目录结构、基础代码与示例输入输出，我负责关键逻辑与测试。",
            "调试与重构：把报错栈与最小复现例给 AI，快速定位可能原因，再由我确认与修复。",
            "质量保障：对关键模块强制补单元测试/验收用例；对外部依赖做版本固定，避免“能跑但不稳”。",
        ],
    )
    _h2(doc, "4. 风险与规范：让赋能可持续")
    _numbered(
        doc,
        [
            "隐私与安全：不把敏感信息（个人隐私、账号密钥、未公开数据）直接输入第三方模型；必要时脱敏或本地处理。",
            "学术诚信：明确哪些部分使用了 AI 辅助（如润色/生成大纲/代码建议），并确保最终内容可追溯、可解释、可复现。",
            "防止过度依赖：对重要结论坚持“能手算/能复述/能复现”，把 AI 当助教而不是替代品。",
        ],
    )

    doc.save(path)


def build_doc_2_notes(path: str) -> None:
    doc = Document()
    _set_default_font(doc, "宋体", 12)

    _title(doc, "人工智能课外学习笔记与实验记录", "（作业 2：分专题笔记 + 网络资源 + 书单 + ≥5算法实验）")
    _meta_block(doc)

    _h1(doc, "一、网络学习资源使用笔记（学习通 / 慕课网 / bilibili / CSDN / 知乎）")
    _h2(doc, "1. 学习通（超星）")
    _bullets(
        doc,
        [
            "定位：课程学习、作业测验、资料下载与班级互动。",
            "使用方法：按章节学习→做章节测验→用错题本复盘→把知识点整理到个人笔记。",
            "收获：适合建立系统框架与完成课程要求，缺点是实践部分需要自补。",
        ],
    )
    _h2(doc, "2. 慕课网（imooc）/ 中国大学MOOC")
    _bullets(
        doc,
        [
            "定位：体系化课程，适合按项目或技能树学习（Python、数据分析、深度学习入门等）。",
            "使用方法：跟随课程代码→自己复现一遍→改造为小项目（加入数据/评估/可视化）。",
            "收获：工程化讲解更强，便于“学完就能做”。",
        ],
    )
    _h2(doc, "3. bilibili")
    _bullets(
        doc,
        [
            "定位：高频更新的公开视频，适合查补短板（如线代、概率、PyTorch 实战）。",
            "使用方法：挑选高质量系列课→边看边敲代码→暂停做推导与总结→整理成专题笔记。",
            "收获：节奏快、例子多，但需要自己做结构化整理，避免碎片化。",
        ],
    )
    _h2(doc, "4. CSDN")
    _bullets(
        doc,
        [
            "定位：踩坑与工程经验聚合（环境配置、报错解决、实现细节）。",
            "使用方法：遇到问题先看官方文档/报错栈→再用 CSDN 查解决方案→最终把可行步骤写进自己的“踩坑笔记”。",
            "注意：内容质量参差，必须交叉验证，避免复制粘贴导致理解缺失。",
        ],
    )
    _h2(doc, "5. 知乎")
    _bullets(
        doc,
        [
            "定位：概念解释、行业讨论、学习路线与书单建议。",
            "使用方法：用知乎获取“问题背景+不同观点”→再回到论文/书籍验证→形成自己的结论。",
            "注意：避免把观点当事实；优先看引用充分、可复现的回答。",
        ],
    )

    _h2(doc, "6. WaytoAGI（飞书知识库学习路线导航）")
    _para(
        doc,
        "补充资源：WaytoAGI 是面向“大模型/智能体”方向的学习路线与资源导航型知识库，适合在入门期快速建立路线、在实践期按模块查资料。"
    )
    _bullets(
        doc,
        [
            "推荐用法：先看总体路线→选定一个模块（Prompt/RAG/Agent/评测/部署）→按清单做小实验→把结果写成可复现笔记。",
            "优势：资源聚合度高，便于查找课程/论文/工具；适合做“从 0 到 1”的学习路径规划。",
            "注意：清单型资料容易“收藏不学习”，需要设定每周产出（代码/复盘/总结）来闭环。",
            "链接（作业引用）：https://waytoagi.feishu.cn/wiki/PFXnwBTsEiGwGGk2QQFcdTWrnlb?table=blkjooAlLFNtvKJ2",
        ],
    )

    _h1(doc, "二、人工智能学习相关图书/资料清单（个人推荐）")
    _bullets(
        doc,
        [
            "《机器学习》（周志华）：经典入门，适合建立整体框架（西瓜书）。",
            "《统计学习方法》（李航）：偏理论与推导，适合理解经典模型与优化思想。",
            "《深度学习》（Goodfellow 等）：深度学习基础理论与方法论。",
            "《Reinforcement Learning》（Sutton & Barto）：强化学习经典教材（进阶）。",
            "《Pattern Recognition and Machine Learning》（Bishop）：概率图模型与贝叶斯视角（进阶）。",
            "官方文档与课程：PyTorch 官方教程、CS231n/CS224n、Stanford/DeepLearning.AI 等。",
        ],
    )

    _h1(doc, "三、专题学习笔记与实验记录（算法数量 ≥ 5）")
    _para(
        doc,
        "说明：下列实验记录以“可复现的小实验/玩具数据”为主，强调过程记录、关键参数与反思。可根据课程要求替换为自己真实跑过的数据与截图。"
    )

    # --- 专题 0：WaytoAGI 学习路线速记（补充） ---
    _h2(doc, "专题 0：WaytoAGI 学习路线速记（补充）")
    _para(
        doc,
        "我参考 WaytoAGI 的路线型资料，把“大模型应用学习”整理为一个可执行的路线：基础 → 提示词 → 检索增强 → 智能体 → 评测与安全 → 部署与迭代。"
        "这样做的目的，是把零散资源变成“每一步都有产出”的行动清单。"
    )
    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "阶段"
    table.cell(0, 1).text = "学习重点（做什么）"
    table.cell(0, 2).text = "产出（交付物）"
    plan_rows = [
        ("基础", "补齐概率/线代/优化与 Python 工程能力", "一份数学/代码速查笔记 + 1 个小数据实验"),
        ("Prompt", "结构化提示词、Few-shot、输出格式约束", "Prompt 模板库 + 20 条对比实验记录"),
        ("RAG", "切分/召回/重排/引用，降低幻觉", "一个可跑通的 RAG Demo + 评测表"),
        ("Agent", "任务分解、工具调用、记忆与反思", "一个“能用工具完成任务”的智能体脚本"),
        ("评测/安全", "幻觉、偏见、隐私、鲁棒性与提示注入", "一套评测用例 + 风险清单与改进措施"),
        ("部署/迭代", "日志、监控、成本、延迟与版本管理", "可复用的项目结构 + 部署说明"),
    ]
    for r, (a, b, c) in enumerate(plan_rows, start=1):
        table.cell(r, 0).text = a
        table.cell(r, 1).text = b
        table.cell(r, 2).text = c
    _para(doc, "引用链接（WaytoAGI）：https://waytoagi.feishu.cn/wiki/PFXnwBTsEiGwGGk2QQFcdTWrnlb?table=blkjooAlLFNtvKJ2")

    # --- 专题 1：不确定性推理（贝叶斯） ---
    _h2(doc, "专题 1：不确定性推理——贝叶斯公式与朴素贝叶斯")
    _para(doc, "核心思想：在不确定信息下，用概率表达信念，并根据新证据更新。")
    _para(doc, "贝叶斯公式：P(A|B) = P(B|A)·P(A) / P(B)")
    _code_block(
        doc,
        "Naive Bayes (伪代码)\n"
        "输入：训练集 D，类别集合 C\n"
        "1) 估计先验 P(c)\n"
        "2) 估计条件概率 P(x_i | c)（离散用计数+拉普拉斯平滑）\n"
        "3) 预测：c* = argmax_c P(c) * Π_i P(x_i|c)\n",
    )
    _h2(doc, "实验记录 1：朴素贝叶斯做垃圾邮件/文本分类（示例）")
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("数据", "公开短信垃圾/评论数据（或自建小语料）"),
        ("特征", "词袋/TF-IDF；文本清洗：分词、去停用词"),
        ("模型", "Multinomial Naive Bayes + 拉普拉斯平滑"),
        ("评估", "Accuracy、Precision、Recall、F1；混淆矩阵"),
        ("结果（示例）", "F1 在小数据上通常稳定；对高频词敏感"),
        ("反思", "独立性假设过强；可用 n-gram 或逻辑回归/深度模型提升"),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    # --- 专题 2：确定性推理（规则/前向后向链） ---
    _h2(doc, "专题 2：确定性推理——规则系统、前向链与后向链")
    _para(
        doc,
        "核心思想：用“事实 + 规则”表达知识，通过逻辑匹配得到可确定的结论。适合流程明确、规则可枚举的领域（故障诊断、流程审批等）。"
    )
    _bullets(doc, ["前向链（数据驱动）：从已知事实出发，不断触发规则推新事实。", "后向链（目标驱动）：从目标结论出发，反推需要满足的前提。"])
    _code_block(
        doc,
        "Forward Chaining (简化伪代码)\n"
        "事实集 F 初始化为已知事实\n"
        "while 存在规则 r: (前提 ⊆ F 且 结论 ∉ F):\n"
        "    F = F ∪ {r.结论}\n"
        "输出：扩展后的事实集 F\n",
    )
    _h2(doc, "实验记录 2：规则推理做“电脑无法上网”故障诊断（示例）")
    _bullets(
        doc,
        [
            "事实：网卡已启用、WiFi 连接正常/异常、能否 ping 网关、能否解析 DNS 等。",
            "规则示例：若“能 ping 网关但不能解析域名”→ 推断“DNS 配置问题”。",
            "结果：能快速定位常见问题；规则越多维护成本越高，需要版本管理与冲突检测。",
        ],
    )

    # --- 专题 3：计算智能（遗传算法） ---
    _h2(doc, "专题 3：计算智能——遗传算法（Genetic Algorithm, GA）")
    _para(doc, "核心思想：用“选择-交叉-变异”在解空间中搜索，适合组合优化与非凸问题。")
    _code_block(
        doc,
        "GA (伪代码)\n"
        "初始化种群 P\n"
        "repeat 直到满足终止条件:\n"
        "  1) 计算适应度 fitness\n"
        "  2) 选择（轮盘赌/锦标赛）得到父代\n"
        "  3) 交叉产生子代\n"
        "  4) 变异增加多样性\n"
        "  5) 精英保留（可选）\n"
        "输出：最优个体\n",
    )
    _h2(doc, "实验记录 3：GA 解旅行商问题（TSP，示例）")
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    rows = [
        ("问题", "给定 N 个城市坐标，求最短巡回路径"),
        ("编码", "排列编码（城市访问顺序）"),
        ("适应度", "fitness = 1 / (路径长度 + ε)"),
        ("交叉/变异", "部分映射交叉 PMX；交换变异 swap"),
        ("参数（示例）", "种群=100，迭代=300，交叉率=0.9，变异率=0.05"),
        ("结果（示例）", "路径长度随迭代下降，后期趋于稳定；精英保留可避免退化"),
        ("反思", "易早熟收敛；可增加变异/使用多种交叉策略/与局部搜索结合"),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    # --- 专题 4：计算智能（粒子群 PSO） ---
    _h2(doc, "专题 4：计算智能——粒子群优化（Particle Swarm Optimization, PSO）")
    _para(doc, "核心思想：粒子在速度-位置更新中向个体最优与全局最优靠拢，适合连续优化。")
    _code_block(
        doc,
        "PSO 更新公式（示意）\n"
        "v = w*v + c1*r1*(p_best - x) + c2*r2*(g_best - x)\n"
        "x = x + v\n",
    )
    _h2(doc, "实验记录 4：PSO 优化二元函数（Rastrigin/ Sphere，示例）")
    _bullets(
        doc,
        [
            "目标：最小化 f(x,y)，观察收敛速度与稳定性。",
            "设置：粒子数=30，迭代=200，惯性权重 w 从 0.9 线性降到 0.4。",
            "结果（示例）：Sphere 函数收敛快；Rastrigin 可能陷入局部，需要参数调节或多次运行取最优。",
            "反思：PSO 对参数敏感；可加入速度限制与随机重启提升鲁棒性。",
        ],
    )

    # --- 专题 5：计算智能（蚁群 ACO） ---
    _h2(doc, "专题 5：计算智能——蚁群算法（Ant Colony Optimization, ACO）")
    _para(doc, "核心思想：通过信息素与启发式因子引导搜索，适合路径规划与组合优化。")
    _code_block(
        doc,
        "ACO（核心要点）\n"
        "1) 蚂蚁按概率选择下一节点：与信息素 τ 和启发式 η（如 1/距离）相关\n"
        "2) 完成路径后更新信息素：蒸发 + 增强（优路径增益更大）\n",
    )
    _h2(doc, "实验记录 5：ACO 做最短路径/小规模 TSP（示例）")
    _bullets(
        doc,
        [
            "数据：构造 15~30 个点的 TSP 或网格地图。",
            "观察：信息素逐步集中到较优边；蒸发率过低会早熟，过高则难收敛。",
            "结论：在组合优化问题上直观好用，但时间复杂度较高，需要控制规模与参数。",
        ],
    )

    # --- 专题 6：专家系统 ---
    _h2(doc, "专题 6：专家系统（Expert System）")
    _para(
        doc,
        "核心组成：知识库（规则/事实）、推理机（前向/后向链）、解释模块（为什么/怎么得到的）、用户接口。"
        "专家系统的价值在于：把专家经验显式化，便于复用与审计。"
    )
    _h2(doc, "实验记录 6：小型“课程选课建议”专家系统（示例）")
    _bullets(
        doc,
        [
            "输入：基础（数学/编程）、兴趣方向（CV/NLP/数据）、时间预算。",
            "规则：若“数学较弱且想学深度学习”→ 推荐先补概率/线代；若“时间少”→ 推荐项目驱动课程。",
            "输出：课程建议 + 解释（触发了哪些规则）。",
        ],
    )

    # --- 专题 7：深度学习 ---
    _h2(doc, "专题 7：深度学习——从 MLP/CNN 到 Transformer")
    _para(doc, "核心思想：通过多层网络学习表示（Representation），用反向传播与梯度下降优化参数。")
    _bullets(
        doc,
        [
            "MLP：适合表格/简单特征，易作为 baseline。",
            "CNN：擅长图像/局部模式提取，利用卷积与池化减少参数、增强平移不变性。",
            "Transformer：通过自注意力建模全局依赖，成为 NLP/多模态的核心架构。",
        ],
    )
    _h2(doc, "实验记录 7：CNN 进行手写数字识别（示例）")
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    rows = [
        ("数据", "MNIST（或自建简单图片分类）"),
        ("网络", "Conv-ReLU-Pool ×2 + FC + Softmax"),
        ("训练", "epoch=5~10，batch=64，优化器 Adam，学习率 1e-3"),
        ("正则化", "Dropout / 数据增强（可选）"),
        ("评估", "Accuracy；查看误分类样本"),
        ("结果（示例）", "通常可达 98%+；误差集中在相似数字（如 4/9）"),
        ("反思", "数据增强与学习率策略能提升泛化；需注意过拟合与训练稳定性"),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    # --- 专题 8：大模型 ---
    _h2(doc, "专题 8：大模型（LLM）学习笔记：Prompt、RAG 与微调（概念+实践要点）")
    _bullets(
        doc,
        [
            "Prompt：明确角色、目标、约束、输出格式；提供示例（Few-shot）可显著提升稳定性。",
            "RAG：检索增强生成，用“外部知识库+引用”降低幻觉；关键在于切分、召回与重排。",
            "微调：用领域数据对齐风格/任务；小数据优先考虑 LoRA/参数高效微调。",
            "评测：不仅看主观感受，还要有可重复指标（准确率、引用命中、人工抽检规则）。",
        ],
    )
    _h2(doc, "实验记录 8：用 RAG 做“课程知识问答”小系统（示例）")
    _numbered(
        doc,
        [
            "整理资料：课件/笔记/教材章节 → 转为纯文本。",
            "切分：按段落或固定长度切 chunk，保留标题与来源信息。",
            "向量化：用 embedding 模型生成向量并入库（FAISS/向量数据库）。",
            "检索+生成：问题 → 召回 top-k → 组装提示词 → 大模型生成答案并引用来源。",
            "评测：抽 30 个问题，统计“回答正确/引用正确/是否胡编”。",
        ],
    )
    _bullets(doc, ["收获：比纯对话更可靠；局限：资料质量与切分策略决定上限。"])

    _h1(doc, "四、个人网络发布链接（如有则填写；没有可留空或删除本节）")
    _para(doc, "说明：此处仅用于附上自己真实发布内容的链接，请不要填写非本人链接。")
    _bullets(
        doc,
        [
            "朋友圈/公众号：______________________________",
            "CSDN：______________________________",
            "知乎：______________________________",
            "GitHub/Gitee（可选）：______________________________",
        ],
    )

    _h1(doc, "五、附：个人学习笔记模板（便于后续持续记录）")
    _code_block(
        doc,
        "【专题名称】\n"
        "1) 概念定义：\n"
        "2) 关键公式/图示：\n"
        "3) 伪代码/流程：\n"
        "4) 实验记录（数据/参数/结果）：\n"
        "5) 踩坑与解决：\n"
        "6) 反思与改进：\n"
        "7) 扩展阅读链接：\n",
    )

    doc.save(path)


def _auto_pick_output(preferred: str, marker_text: str) -> str:
    """
    用户可能手动改过文件名：优先选用当前目录下“已存在且包含 marker_text”的 docx 文件名；
    若未找到，再退回到 preferred。
    """
    here = Path(".")
    for p in sorted(here.glob("*.docx"), key=lambda x: x.stat().st_mtime, reverse=True):
        try:
            doc = Document(str(p))
        except Exception:
            continue
        head = "\n".join((pp.text or "").strip() for pp in doc.paragraphs[:8])
        if marker_text in head:
            return p.name
    return preferred


def main() -> None:
    parser = argparse.ArgumentParser(description="生成两份 AI 作业 .docx（支持自动适配你重命名后的文件名）")
    parser.add_argument("--doc1", default="", help="作业1输出文件名（可选；留空则自动识别或使用默认名）")
    parser.add_argument("--doc2", default="", help="作业2输出文件名（可选；留空则自动识别或使用默认名）")
    parser.add_argument("--only", choices=["all", "doc1", "doc2"], default="all", help="仅生成指定文档（默认 all）")
    args = parser.parse_args()

    # 默认名：优先使用你当前目录里常见的重命名版本；若不存在，会自动回退。
    default_doc1 = "人工智能个人理解与感悟.docx"
    default_doc2 = "人工智能课外学习笔记与实验记录.docx"
    if not Path(default_doc1).exists():
        default_doc1 = "作业_01_人工智能个人理解与感悟_23050342008_高榆展.docx"
    if not Path(default_doc2).exists():
        default_doc2 = "作业_02_人工智能课外学习笔记与实验记录_23050342008_高榆展.docx"

    doc1 = args.doc1.strip() or _auto_pick_output(default_doc1, "人工智能：个人理解与学习感悟")
    doc2 = args.doc2.strip() or _auto_pick_output(default_doc2, "人工智能课外学习笔记与实验记录")

    if args.only in ("all", "doc1"):
        build_doc_1_understanding(doc1)
    if args.only in ("all", "doc2"):
        build_doc_2_notes(doc2)

    print("已生成：")
    if args.only in ("all", "doc1"):
        print(f"- {doc1}")
    if args.only in ("all", "doc2"):
        print(f"- {doc2}")


if __name__ == "__main__":
    main()


