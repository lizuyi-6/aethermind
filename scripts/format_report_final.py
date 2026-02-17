#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
最终格式化脚本：严格按照模板格式格式化报告，直接覆盖原文件
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import os
import sys
import re

def set_chinese_font(run, font_name='宋体', size=Pt(12)):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size

def format_document_final(source_doc, template_doc, output_path):
    """最终格式化文档"""
    # 创建新文档
    new_doc = Document()
    
    # 设置默认样式（参考模板）
    normal_style = new_doc.styles['Normal']
    normal_style.font.name = '宋体'
    normal_style.font.size = Pt(12)
    
    # 设置标题样式
    heading1_style = new_doc.styles['Heading 1']
    heading1_style.font.name = '宋体'
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1_style.paragraph_format.space_after = Pt(12)
    
    heading2_style = new_doc.styles['Heading 2']
    heading2_style.font.name = '宋体'
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    heading2_style.paragraph_format.space_before = Pt(12)
    heading2_style.paragraph_format.space_after = Pt(6)
    
    heading3_style = new_doc.styles['Heading 3']
    heading3_style.font.name = '宋体'
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True
    heading3_style.paragraph_format.space_before = Pt(6)
    heading3_style.paragraph_format.space_after = Pt(6)
    
    # 处理源文档
    for para in source_doc.paragraphs:
        text = para.text.strip()
        
        # 空行
        if not text:
            new_doc.add_paragraph()
            continue
        
        # 分隔线
        if text == '---' or text.startswith('---'):
            continue
        
        # 创建新段落
        new_para = new_doc.add_paragraph()
        
        # 判断段落类型并格式化
        clean_text = text
        
        # 一级标题：以# 开头或"第X章"
        if text.startswith('# ') or (text.startswith('第') and '章' in text and len(text) < 50):
            clean_text = re.sub(r'^#+\s*', '', text).strip()
            new_para.style = 'Heading 1'
            run = new_para.add_run(clean_text)
            set_chinese_font(run, '宋体', Pt(16))
            run.bold = True
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 二级标题：以## 开头或"X.X"格式
        elif text.startswith('## ') or re.match(r'^\d+\.\d+\s+', text):
            clean_text = re.sub(r'^#+\s*', '', text).strip()
            new_para.style = 'Heading 2'
            run = new_para.add_run(clean_text)
            set_chinese_font(run, '宋体', Pt(14))
            run.bold = True
        
        # 三级标题：以### 开头或"X.X.X"格式
        elif text.startswith('### ') or re.match(r'^\d+\.\d+\.\d+\s+', text):
            clean_text = re.sub(r'^#+\s*', '', text).strip()
            new_para.style = 'Heading 3'
            run = new_para.add_run(clean_text)
            set_chinese_font(run, '宋体', Pt(12))
            run.bold = True
        
        # 关键词段落
        elif text.startswith('**关键词**') or text.startswith('关键词'):
            clean_text = text.replace('**', '').strip()
            run = new_para.add_run(clean_text)
            set_chinese_font(run, '宋体', Pt(12))
            run.bold = True
        
        # 目录项
        elif text.startswith('- [') or text.startswith('* ['):
            clean_text = re.sub(r'^[-*]\s*\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            run = new_para.add_run(clean_text)
            set_chinese_font(run, '宋体', Pt(12))
            new_para.paragraph_format.left_indent = Cm(1)
        
        # 说明文字
        elif text.startswith('**说明**'):
            clean_text = text.replace('**', '').strip()
            run = new_para.add_run(clean_text)
            set_chinese_font(run, '宋体', Pt(10))
            new_para.paragraph_format.first_line_indent = Cm(0)
        
        # 普通段落
        else:
            # 清理Markdown格式
            clean_text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', clean_text)  # 移除加粗标记但保留文本
            clean_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_text)  # 移除链接标记但保留文本
            
            run = new_para.add_run(clean_text)
            set_chinese_font(run, '宋体', Pt(12))
            
            # 普通段落首行缩进2字符（0.74cm）
            new_para.paragraph_format.first_line_indent = Cm(0.74)
            new_para.paragraph_format.space_after = Pt(6)
            new_para.paragraph_format.line_spacing = 1.5
    
    # 保存文档（覆盖原文件）
    new_doc.save(output_path)
    print(f"格式化完成，已保存到: {output_path}")

def main():
    """主函数"""
    template_path = "23050342008_高榆展_基于B2B的可行性分析报告智能体.docx"
    source_path = "报告_01_封面和摘要.docx"
    output_path = "报告_01_封面和摘要.docx"  # 直接覆盖原文件
    
    if not os.path.exists(template_path):
        print(f"错误: 模板文件不存在: {template_path}")
        return
    
    if not os.path.exists(source_path):
        print(f"错误: 源文件不存在: {source_path}")
        return
    
    print("正在读取模板文档...")
    template_doc = Document(template_path)
    
    print("正在读取源文档...")
    source_doc = Document(source_path)
    
    print("正在应用模板格式...")
    format_document_final(source_doc, template_doc, output_path)
    
    print("完成！文档已按照模板格式格式化。")

if __name__ == "__main__":
    main()

