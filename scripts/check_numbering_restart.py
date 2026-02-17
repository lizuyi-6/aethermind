from __future__ import annotations

import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def find_target_docx() -> Path:
    """
    避免命令行里直接写中文文件名（在部分终端编码下会乱码）。
    这里通过“文档标题”来定位目标 docx。
    """
    for p in sorted(Path(".").glob("*.docx"), key=lambda x: x.stat().st_mtime, reverse=True):
        try:
            doc = Document(str(p))
        except Exception:
            continue
        head = "\n".join((pp.text or "").strip() for pp in doc.paragraphs[:6])
        if "人工智能：个人理解与学习感悟" in head:
            return p
    raise FileNotFoundError("未找到包含标题“人工智能：个人理解与学习感悟”的 docx。")


def main() -> None:
    docx_path = find_target_docx()
    print("目标文件：", docx_path.name)

    with zipfile.ZipFile(docx_path, "r") as z:
        xml = z.read("word/numbering.xml")
    root = ET.fromstring(xml)

    nums = root.findall(".//w:num", NS)
    starts = root.findall(".//w:startOverride", NS)
    print("numbering.xml 统计：num 总数 =", len(nums), "startOverride 总数 =", len(starts))

    # 再统计文档中 “List Number” 段落使用了多少个不同 numId（越多越说明每段列表都在重置）
    doc = Document(str(docx_path))
    num_ids = set()
    list_number_paras = 0
    for p in doc.paragraphs:
        if getattr(p.style, "name", "") != "List Number":
            continue
        list_number_paras += 1
        # 读 w:numPr/w:numId
        vals = p._p.xpath(".//w:numPr/w:numId/@w:val")
        if vals:
            num_ids.add(vals[0])
    print("文档统计：List Number 段落数 =", list_number_paras, "不同 numId 数 =", len(num_ids))


if __name__ == "__main__":
    main()


